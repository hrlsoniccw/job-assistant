import os
import re
from typing import Optional, Dict, List, Any
from dataclasses import dataclass, field
from datetime import datetime


@dataclass
class ResumeData:
    """结构化简历数据"""
    # 基本信息
    name: str = ""
    phone: str = ""
    email: str = ""
    location: str = ""
    blog: str = ""
    github: str = ""
    
    # 求职意向
    job_title: str = ""
    expected_salary: str = ""
    expected_city: str = ""
    job_type: str = ""  # 全职/兼职/实习
    
    # 工作经历
    work_experience: List[Dict] = field(default_factory=list)
    
    # 项目经历
    project_experience: List[Dict] = field(default_factory=list)
    
    # 教育背景
    education: List[Dict] = field(default_factory=list)
    
    # 技能
    skills: List[str] = field(default_factory=list)
    
    # 证书
    certificates: List[str] = field(default_factory=list)
    
    # 奖项
    awards: List[str] = field(default_factory=list)
    
    # 自我介绍
    self_introduction: str = ""
    
    # 原始文本
    raw_text: str = ""


try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pytesseract
except ImportError:
    pytesseract = None


def allowed_file(filename: str) -> bool:
    """检查文件扩展名是否允许"""
    if '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower()
    return ext in {'txt', 'pdf', 'docx', 'doc', 'jpg', 'jpeg', 'png'}


def get_file_extension(filename: str) -> str:
    """获取文件扩展名"""
    if '.' not in filename:
        return ''
    return filename.rsplit('.', 1)[1].lower()


def parse_resume(file_path: str, filename: str) -> str:
    """解析简历文件，返回纯文本内容"""
    ext = get_file_extension(filename)
    
    try:
        if ext == 'pdf':
            return parse_pdf(file_path)
        elif ext in ('docx', 'doc'):
            return parse_word(file_path)
        elif ext == 'txt':
            return parse_text(file_path)
        elif ext in ('jpg', 'jpeg', 'png'):
            return parse_image(file_path)
        else:
            return parse_text(file_path)
    except Exception as e:
        print(f"解析文件失败 {filename}: {e}")
        return ""


def parse_pdf(file_path: str) -> str:
    """解析PDF文件"""
    if pdfplumber is None:
        raise ImportError("请安装pdfplumber: pip install pdfplumber")
    
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def parse_word(file_path: str) -> str:
    """解析Word文件"""
    if Document is None:
        raise ImportError("请安装python-docx: pip install python-docx")
    
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # 尝试提取表格内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    
    return text


def parse_text(file_path: str) -> str:
    """解析文本文件"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def parse_image(file_path: str) -> str:
    """解析图片文件（OCR识别）"""
    if Image is None:
        raise ImportError("请安装Pillow: pip install Pillow")
    if pytesseract is None:
        raise ImportError("请安装pytesseract: pip install pytesseract")
    
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image, lang='chi_sim+eng')
        return text
    except Exception as e:
        print(f"OCR识别失败: {e}")
        # 如果OCR失败，返回空字符串而不是抛出异常
        return ""


def extract_contact_info(text: str) -> dict:
    """从简历文本中提取联系信息"""
    info = {
        'phone': '',
        'email': '',
        'name': ''
    }
    
    # 提取手机号
    phone_pattern = r'1[3-9]\d{9}'
    phones = re.findall(phone_pattern, text)
    if phones:
        info['phone'] = phones[0]
    
    # 提取邮箱
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    emails = re.findall(email_pattern, text)
    if emails:
        info['email'] = emails[0]
    
    # 提取姓名（假设第一行是姓名）
    lines = text.strip().split('\n')
    if lines:
        info['name'] = lines[0].strip()
    
    return info


def extract_work_experience(text: str) -> Dict[str, Any]:
    """提取工作经历"""
    experiences = []
    projects = []
    education = []
    certificates = []
    awards = []
    self_intro = ""
    
    lines = text.strip().split('\n')
    
    # 工作经历模式
    work_patterns = [
        r'工作经历[：:\-]?',
        r'职业经历[：:\-]?',
        r'任职经历[：:\-]?',
        r'工作经历\s*[-–—]',
        r'(\d{4}[-/]\d{0,2})\s*[-–—]\s*(\d{4}[-/]\d{0,2}|至今|现在)'
    ]
    
    # 项目经历模式
    project_patterns = [
        r'项目经历[：:\-]?',
        r'项目经验[：:\-]?',
        r'项目背景[：:\-]?',
        r'参与项目[：:\-]?'
    ]
    
    # 教育背景模式
    education_patterns = [
        r'教育背景[：:\-]?',
        r'教育经历[：:\-]?',
        r'学术背景[：:\-]?'
    ]
    
    # 证书模式
    certificate_patterns = [
        r'证书资质[：:\-]?',
        r'获得证书[：:\-]?',
        r'资格证书[：:\-]?'
    ]
    
    # 奖项模式
    award_patterns = [
        r'获奖情况[：:\-]?',
        r'荣誉奖励[：:\-]?',
        r'获得奖项[：:\-]?'
    ]
    
    # 自我介绍模式
    intro_patterns = [
        r'自我介绍[：:\-]?',
        r'个人简介[：:\-]?',
        r'关于我[：:\-]?'
    ]
    
    # 当前状态标记
    current_section = 'other'
    section_content = []
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        # 检测章节标题
        is_section_header = False
        
        # 检测工作经历
        for pattern in work_patterns:
            if re.search(pattern, line):
                if section_content and current_section == 'work':
                    experiences.append(extract_company_info('\n'.join(section_content)))
                section_content = []
                current_section = 'work'
                is_section_header = True
                break
        
        if not is_section_header:
            # 检测项目经历
            for pattern in project_patterns:
                if re.search(pattern, line):
                    if section_content and current_section == 'project':
                        projects.append(extract_project_info('\n'.join(section_content)))
                    section_content = []
                    current_section = 'project'
                    is_section_header = True
                    break
        
        if not is_section_header:
            # 检测教育背景
            for pattern in education_patterns:
                if re.search(pattern, line):
                    if section_content and current_section == 'education':
                        education.append(extract_education_info('\n'.join(section_content)))
                    section_content = []
                    current_section = 'education'
                    is_section_header = True
                    break
        
        if not is_section_header:
            # 检测证书
            for pattern in certificate_patterns:
                if re.search(pattern, line):
                    if section_content and current_section == 'certificate':
                        certificates.extend(extract_items('\n'.join(section_content)))
                    section_content = []
                    current_section = 'certificate'
                    is_section_header = True
                    break
        
        if not is_section_header:
            # 检测奖项
            for pattern in award_patterns:
                if re.search(pattern, line):
                    if section_content and current_section == 'award':
                        awards.extend(extract_items('\n'.join(section_content)))
                    section_content = []
                    current_section = 'award'
                    is_section_header = True
                    break
        
        if not is_section_header:
            # 检测自我介绍
            for pattern in intro_patterns:
                if re.search(pattern, line):
                    self_intro = line
                    current_section = 'intro'
                    is_section_header = True
                    break
        
        if not is_section_header:
            section_content.append(line)
    
    # 处理最后一个section
    if section_content:
        if current_section == 'work':
            experiences.append(extract_company_info('\n'.join(section_content)))
        elif current_section == 'project':
            projects.append(extract_project_info('\n'.join(section_content)))
        elif current_section == 'education':
            education.append(extract_education_info('\n'.join(section_content)))
        elif current_section == 'certificate':
            certificates.extend(extract_items('\n'.join(section_content)))
        elif current_section == 'award':
            awards.extend(extract_items('\n'.join(section_content)))
    
    # 如果没有检测到章节，尝试全文搜索
    if not experiences:
        experiences = extract_experiences_from_text(text)
    if not projects:
        projects = extract_projects_from_text(text)
    if not education:
        education = extract_education_from_text(text)
    
    return {
        'work': experiences,
        'projects': projects,
        'education': education,
        'certificates': certificates,
        'awards': awards,
        'self_introduction': self_intro
    }


def extract_company_info(text: str) -> Dict:
    """提取公司信息"""
    info = {
        'company': '',
        'position': '',
        'start_date': '',
        'end_date': '',
        'description': '',
        'achievements': []
    }
    
    lines = text.strip().split('\n')
    
    # 提取公司名称（通常是第一行或包含特定关键词）
    company_keywords = ['公司', '科技', '网络', '信息', '有限', '集团', '企业']
    for line in lines:
        for keyword in company_keywords:
            if keyword in line and len(line) < 50:
                info['company'] = line.strip()
                break
        if info['company']:
            break
    
    # 提取职位
    position_keywords = ['工程师', '经理', '主管', '总监', '架构师', '专家', '专员', '分析师', '顾问']
    for line in lines:
        for keyword in position_keywords:
            if keyword in line:
                info['position'] = line.strip()
                break
        if info['position']:
            break
    
    # 提取时间
    date_pattern = r'(\d{4}[-/]\d{0,2})\s*[-–—至]\s*(\d{4}[-/]\d{0,2}|至今|现在)'
    dates = re.findall(date_pattern, text)
    if dates:
        info['start_date'] = dates[0][0]
        info['end_date'] = dates[0][1]
    
    # 提取描述
    info['description'] = text
    
    # 提取成就（量化）
    achievement_patterns = [
        r'提升了?(\d+%)',
        r'降低了?(\d+%)',
        r'节省了?(\d+[\w]+)',
        r'优化了?\s*(\w+)',
        r'完成[了]?(\w+)',
        r'实现[了]?(\w+)'
    ]
    
    achievements = []
    for line in lines:
        for pattern in achievement_patterns:
            matches = re.findall(pattern, line)
            if matches:
                achievements.append(line.strip())
                break
    
    info['achievements'] = achievements
    
    return info


def extract_project_info(text: str) -> Dict:
    """提取项目信息"""
    info = {
        'name': '',
        'role': '',
        'start_date': '',
        'end_date': '',
        'description': '',
        'tech_stack': []
    }
    
    lines = text.strip().split('\n')
    
    # 提取项目名称（通常第一行）
    if lines:
        info['name'] = lines[0].strip()
    
    # 提取角色
    role_keywords = ['负责人', '主导', '参与', '核心成员', '开发', '设计']
    for line in lines:
        for keyword in role_keywords:
            if keyword in line:
                info['role'] = line.strip()
                break
        if info['role']:
            break
    
    # 提取时间
    date_pattern = r'(\d{4}[-/]\d{0,2})\s*[-–—至]\s*(\d{4}[-/]\d{0,2}|至今|现在)'
    dates = re.findall(date_pattern, text)
    if dates:
        info['start_date'] = dates[0][0]
        info['end_date'] = dates[0][1]
    
    # 提取技术栈
    tech_keywords = ['Python', 'Java', 'React', 'Vue', 'MySQL', 'Redis', 'Docker', 'AWS']
    for line in lines:
        for tech in tech_keywords:
            if tech in line and tech not in info['tech_stack']:
                info['tech_stack'].append(tech)
    
    info['description'] = text
    
    return info


def extract_education_info(text: str) -> Dict:
    """提取教育背景"""
    info = {
        'school': '',
        'degree': '',
        'major': '',
        'start_date': '',
        'end_date': '',
        'description': ''
    }
    
    lines = text.strip().split('\n')
    
    # 学校
    school_keywords = ['大学', '学院', '研究生', '本科', '博士', '硕士']
    for line in lines:
        for keyword in school_keywords:
            if keyword in line:
                info['school'] = line.strip()
                break
        if info['school']:
            break
    
    # 学位
    degree_patterns = ['博士', '硕士', '本科', '大专', 'MBA', 'EMBA']
    for line in lines:
        for degree in degree_patterns:
            if degree in line:
                info['degree'] = degree
                break
    
    # 专业
    major_keywords = ['计算机', '软件工程', '电子信息', '机械', '自动化', '数学', '物理', '经济', '管理']
    for line in lines:
        for major in major_keywords:
            if major in line:
                info['major'] = line.strip()
                break
    
    # 时间
    date_pattern = r'(\d{4}[-/]\d{0,2})\s*[-–—至]\s*(\d{4}[-/]\d{0,2})'
    dates = re.findall(date_pattern, text)
    if dates:
        info['start_date'] = dates[0][0]
        info['end_date'] = dates[0][1]
    
    info['description'] = text
    
    return info


def extract_items(text: str) -> List[str]:
    """提取项目列表"""
    items = []
    lines = text.strip().split('\n')
    
    for line in lines:
        line = line.strip()
        if line and len(line) > 2:
            # 去除前缀符号
            cleaned = re.sub(r'^[\d\.\、\•\-\*]\s*', '', line)
            if cleaned:
                items.append(cleaned)
    
    return items


def extract_experiences_from_text(text: str) -> List[Dict]:
    """从全文提取工作经历"""
    experiences = []
    
    # 查找包含公司名的段落
    company_pattern = r'([^\n]{5,30}公司[^\n]*)'
    companies = re.findall(company_pattern, text)
    
    for company in companies[:5]:  # 最多5个
        exp = {
            'company': company.strip(),
            'position': '',
            'start_date': '',
            'end_date': '',
            'description': '',
            'achievements': []
        }
        experiences.append(exp)
    
    return experiences


def extract_projects_from_text(text: str) -> List[Dict]:
    """从全文提取项目经历"""
    projects = []
    
    # 查找项目关键词
    project_keywords = ['项目', '系统', '平台', '产品', '网站', 'APP']
    
    for keyword in project_keywords:
        pattern = rf'[^\n]*{keyword}[^\n]{{5,50}}'
        matches = re.findall(pattern, text)
        for match in matches[:3]:
            proj = {
                'name': match.strip()[:30],
                'role': '',
                'start_date': '',
                'end_date': '',
                'description': match.strip(),
                'tech_stack': []
            }
            projects.append(proj)
    
    return projects


def extract_education_from_text(text: str) -> List[Dict]:
    """从全文提取教育背景"""
    education = []
    
    # 查找大学
    university_pattern = r'([^\n]{5,30}(?:大学|学院)[^\n]*)'
    universities = re.findall(university_pattern, text)
    
    for uni in universities[:3]:
        edu = {
            'school': uni.strip(),
            'degree': '',
            'major': '',
            'start_date': '',
            'end_date': '',
            'description': uni.strip()
        }
        education.append(edu)
    
    return education


def extract_skills(text: str) -> list:
    """从简历中提取技能关键词"""
    skills = []
    
    # 扩展技术关键词库
    tech_keywords = [
        # 编程语言
        'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'Go', 'Rust', 
        'Ruby', 'PHP', 'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB', 'SQL', 'HTML', 'CSS',
        
        # 前端框架
        'React', 'Vue', 'Vue3', 'Angular', 'Next.js', 'Nuxt.js', 'Svelte', 'jQuery',
        'Bootstrap', 'Tailwind CSS', 'Ant Design', 'Element UI', 'Material-UI',
        
        # 后端框架
        'Django', 'Flask', 'FastAPI', 'Spring', 'Spring Boot', 'Spring Cloud',
        'Express', 'NestJS', 'Koa', 'Gin', 'Echo', 'Laravel', 'Ruby on Rails',
        
        # 数据库
        'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Oracle', 'SQL Server',
        'Elasticsearch', 'ClickHouse', 'HBase', 'Cassandra', 'DynamoDB',
        
        # 云和DevOps
        'AWS', 'Azure', 'GCP', 'Aliyun', 'TencentCloud', 'Docker', 'Kubernetes',
        'Jenkins', 'GitLab CI', 'GitHub Actions', 'Travis CI', 'Ansible', 'Terraform',
        'Linux', 'Nginx', 'Apache', 'Tomcat', 'Zookeeper', 'Kafka', 'RabbitMQ',
        
        # 数据分析
        'Pandas', 'NumPy', 'Spark', 'Hadoop', 'Hive', 'Flink', 'Storm',
        'Tableau', 'Power BI', 'Excel', 'SPSS', 'FineReport',
        
        # AI/ML
        'TensorFlow', 'PyTorch', 'Keras', 'Scikit-learn', 'XGBoost', 'LightGBM',
        'OpenCV', 'NLP', 'Natural Language Processing', 'Computer Vision',
        'Speech Recognition', 'OCR', 'Transformer', 'BERT', 'GPT',
        
        # 大模型
        'LLM', 'Large Language Model', 'ChatGPT', 'Claude', '文心一言',
        'LangChain', 'Prompt Engineering', 'RAG', '向量数据库',
        
        # 测试
        'Selenium', 'Appium', 'JUnit', 'Pytest', 'Postman', 'JMeter',
        
        # 产品和设计
        'Axure', 'Figma', 'Sketch', 'Adobe XD', 'Photoshop', 'Illustrator',
        '产品设计', '需求分析', '用户研究', '竞品分析', 'PRD', 'MRD',
        
        # 软技能
        '项目管理', '团队协作', '沟通能力', '解决问题', '学习能力',
        'Agile', 'Scrum', 'Kanban', '敏捷开发',
        
        # 其他技术
        'RESTful API', 'GraphQL', 'WebSocket', 'Microservices', 'Serverless',
        'Blockchain', 'Solidity', 'Web3', 'Socket.io',
        
        # 通用工具
        'Git', 'SVN', 'Markdown', 'LaTeX', 'Jira', 'Confluence', 'Notion'
    ]
    
    text_lower = text.lower()
    for keyword in tech_keywords:
        if keyword.lower() in text_lower:
            skills.append(keyword)
    
    # 去重并保持顺序
    seen = set()
    unique_skills = []
    for skill in skills:
        if skill not in seen:
            seen.add(skill)
            unique_skills.append(skill)
    
    return unique_skills


def suggest_job_positions(skills: list, work_experience: list) -> list:
    """根据技能和工作经验推荐合适的岗位方向"""
    positions = []
    
    # 技术岗位分类
    tech_positions = {
        '后端开发': ['Python', 'Java', 'Go', 'Spring', 'Django', 'Flask', 'MySQL', 'Redis'],
        '前端开发': ['React', 'Vue', 'HTML', 'CSS', 'JavaScript', 'TypeScript', 'Node.js'],
        '全栈开发': ['Python', 'JavaScript', 'React', 'Node.js', 'MySQL'],
        '数据分析': ['Python', 'Pandas', 'SQL', 'Excel', 'Tableau', 'Spark'],
        '机器学习': ['Python', 'TensorFlow', 'PyTorch', 'Scikit-learn', 'NLP'],
        'DevOps': ['Docker', 'Kubernetes', 'Jenkins', 'AWS', 'Linux', 'Git'],
        '移动开发': ['Swift', 'Kotlin', 'React Native', 'Android', 'iOS'],
        '产品经理': ['Axure', 'Figma', '需求分析', '产品设计', '用户研究']
    }
    
    for position, required_skills in tech_positions.items():
        match_count = sum(1 for skill in skills if skill in required_skills)
        if match_count >= 2:
            positions.append({
                'position': position,
                'match_score': match_count / len(required_skills),
                'matched_skills': [s for s in skills if s in required_skills]
            })
    
    # 按匹配度排序
    positions.sort(key=lambda x: x['match_score'], reverse=True)
    
    return positions[:5]  # 返回前5个推荐岗位


def parse_resume_full(file_path: str, filename: str) -> ResumeData:
    """
    完整解析简历，返回结构化数据
    返回: ResumeData 对象
    """
    # 1. 解析原始文本
    raw_text = parse_resume(file_path, filename)
    
    if not raw_text:
        return ResumeData(raw_text="")
    
    # 2. 创建简历数据对象
    resume = ResumeData(raw_text=raw_text)
    
    # 3. 提取联系信息
    contact_info = extract_contact_info(raw_text)
    resume.name = contact_info.get('name', '')
    resume.phone = contact_info.get('phone', '')
    resume.email = contact_info.get('email', '')
    
    # 4. 提取技能
    resume.skills = extract_skills(raw_text)
    
    # 5. 提取工作经历、项目、教育等
    experience_data = extract_work_experience(raw_text)
    resume.work_experience = experience_data.get('work', [])
    resume.project_experience = experience_data.get('projects', [])
    resume.education = experience_data.get('education', [])
    resume.certificates = experience_data.get('certificates', [])
    resume.awards = experience_data.get('awards', [])
    resume.self_introduction = experience_data.get('self_introduction', '')
    
    # 6. 提取求职意向（如果有）
    job_intention = extract_job_intention(raw_text)
    resume.job_title = job_intention.get('title', '')
    resume.expected_salary = job_intention.get('salary', '')
    resume.expected_city = job_intention.get('city', '')
    resume.job_type = job_intention.get('type', '')
    
    return resume


def extract_job_intention(text: str) -> Dict:
    """提取求职意向"""
    intention = {
        'title': '',
        'salary': '',
        'city': '',
        'type': ''
    }
    
    # 求职意向关键词
    intention_patterns = [
        r'求职意向[：:\-]?',
        r'期望职位[：:\-]?',
        r'期望岗位[：:\-]?',
        r'目标职位[：:\-]?',
        r'意向职位[：:\-]?'
    ]
    
    salary_patterns = [
        r'期望薪资[：:\-]?',
        r'期望薪酬[：:\-]?',
        r'薪资要求[：:\-]?'
    ]
    
    city_patterns = [
        r'期望城市[：:\-]?',
        r'工作地点[：:\-]?',
        r'意向城市[：:\-]?'
    ]
    
    type_patterns = [
        r'工作性质[：:\-]?',
        r'工作类型[：:\-]?',
        r'求职类型[：:\-]?'
    ]
    
    # 提取求职意向
    for pattern in intention_patterns:
        match = re.search(pattern + r'([^\n]+)', text)
        if match:
            intention['title'] = match.group(1).strip()
            break
    
    # 提取薪资
    for pattern in salary_patterns:
        match = re.search(pattern + r'([^\n]+)', text)
        if match:
            intention['salary'] = match.group(1).strip()
            break
    
    # 提取城市
    for pattern in city_patterns:
        match = re.search(pattern + r'([^\n]+)', text)
        if match:
            intention['city'] = match.group(1).strip()
            break
    
    # 提取工作类型
    for pattern in type_patterns:
        match = re.search(pattern + r'([^\n]+)', text)
        if match:
            intention['type'] = match.group(1).strip()
            break
    
    return intention


def resume_to_dict(resume: ResumeData) -> Dict:
    """将ResumeData转换为字典（便于JSON序列化）"""
    return {
        'name': resume.name,
        'phone': resume.phone,
        'email': resume.email,
        'location': resume.location,
        'blog': resume.blog,
        'github': resume.github,
        'job_title': resume.job_title,
        'expected_salary': resume.expected_salary,
        'expected_city': resume.expected_city,
        'job_type': resume.job_type,
        'work_experience': resume.work_experience,
        'project_experience': resume.project_experience,
        'education': resume.education,
        'skills': resume.skills,
        'certificates': resume.certificates,
        'awards': resume.awards,
        'self_introduction': resume.self_introduction,
        'raw_text_preview': resume.raw_text[:500] if resume.raw_text else ''
    }
