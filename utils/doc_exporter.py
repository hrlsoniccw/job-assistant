import os
import html
from typing import Dict, List, Optional
from datetime import datetime


def format_date(date_str: str) -> str:
    """标准化日期格式"""
    if not date_str:
        return ""
    date_str = date_str.strip()
    if date_str in ['至今', '现在', 'Present', 'Current']:
        return "至今"
    try:
        for fmt in ['%Y-%m', '%Y/%m', '%Y.%m', '%Y']:
            try:
                dt = datetime.strptime(date_str, fmt)
                return dt.strftime('%Y.%m')
            except ValueError:
                continue
        return date_str
    except:
        return date_str


def convert_resume_to_dict(resume: object) -> Dict:
    """将ResumeData对象转换为字典"""
    if hasattr(resume, '__dict__'):
        return {
            'name': getattr(resume, 'name', ''),
            'phone': getattr(resume, 'phone', ''),
            'email': getattr(resume, 'email', ''),
            'location': getattr(resume, 'location', ''),
            'blog': getattr(resume, 'blog', ''),
            'github': getattr(resume, 'github', ''),
            'job_title': getattr(resume, 'job_title', ''),
            'expected_salary': getattr(resume, 'expected_salary', ''),
            'expected_city': getattr(resume, 'expected_city', ''),
            'job_type': getattr(resume, 'job_type', ''),
            'work_experience': getattr(resume, 'work_experience', []),
            'project_experience': getattr(resume, 'project_experience', []),
            'education': getattr(resume, 'education', []),
            'skills': getattr(resume, 'skills', []),
            'certificates': getattr(resume, 'certificates', []),
            'awards': getattr(resume, 'awards', []),
            'self_introduction': getattr(resume, 'self_introduction', ''),
            'raw_text': getattr(resume, 'raw_text', '')
        }
    elif isinstance(resume, dict):
        return resume
    else:
        return {}


class WordExporter:
    """Word文档导出器"""
    
    @staticmethod
    def export(resume_data: Dict, output_path: str) -> bool:
        """导出为Word文档"""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            doc.add_heading(resume_data.get('name', ''), 0)
            
            if resume_data.get('job_title'):
                p = doc.add_paragraph(resume_data['job_title'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            contact = []
            if resume_data.get('phone'):
                contact.append(f"电话: {resume_data['phone']}")
            if resume_data.get('email'):
                contact.append(f"邮箱: {resume_data['email']}")
            if resume_data.get('location'):
                contact.append(f"地点: {resume_data['location']}")
            
            if contact:
                p = doc.add_paragraph(" | ".join(contact))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if resume_data.get('self_introduction'):
                doc.add_heading('个人简介', level=1)
                doc.add_paragraph(resume_data['self_introduction'])
            
            if resume_data.get('skills'):
                doc.add_heading('专业技能', level=1)
                doc.add_paragraph(" | ".join(resume_data['skills']))
            
            if resume_data.get('work_experience'):
                doc.add_heading('工作经历', level=1)
                for exp in resume_data['work_experience']:
                    WordExporter._add_work_item(doc, exp)
            
            if resume_data.get('project_experience'):
                doc.add_heading('项目经历', level=1)
                for proj in resume_data['project_experience']:
                    WordExporter._add_project_item(doc, proj)
            
            if resume_data.get('education'):
                doc.add_heading('教育背景', level=1)
                for edu in resume_data['education']:
                    WordExporter._add_education_item(doc, edu)
            
            if resume_data.get('certificates'):
                doc.add_heading('证书资质', level=1)
                for cert in resume_data['certificates']:
                    doc.add_paragraph(f"• {cert}")
            
            if resume_data.get('awards'):
                doc.add_heading('荣誉奖励', level=1)
                for award in resume_data['awards']:
                    doc.add_paragraph(f"• {award}")
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Word导出失败: {e}")
            return False
    
    @staticmethod
    def export_to_bytes(resume_data: Dict) -> bytes:
        """导出为Word字节流"""
        import io
        import tempfile
        import os
        
        buffer = io.BytesIO()
        temp_path = tempfile.mktemp(suffix='.docx')
        
        try:
            WordExporter.export(resume_data, temp_path)
            with open(temp_path, 'rb') as f:
                buffer.write(f.read())
            return buffer.getvalue()
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    @staticmethod
    def _add_work_item(doc, exp: Dict):
        """添加工作经历条目"""
        if exp.get('company'):
            p = doc.add_paragraph()
            p.add_run(exp['company']).bold = True
            if exp.get('position'):
                p.add_run(f" | {exp['position']}")
            if exp.get('start_date') or exp.get('end_date'):
                date = f"{format_date(exp.get('start_date', ''))} - {format_date(exp.get('end_date', ''))}"
                p.add_run(f" ({date})")
        
        if exp.get('description'):
            doc.add_paragraph(exp['description'].strip())
        
        if exp.get('achievements'):
            for ach in exp['achievements'][:3]:
                doc.add_paragraph(f"◆ {ach}")
        
        doc.add_paragraph()
    
    @staticmethod
    def _add_project_item(doc, proj: Dict):
        """添加项目经历条目"""
        if proj.get('name'):
            p = doc.add_paragraph()
            p.add_run(proj['name']).bold = True
            if proj.get('role'):
                p.add_run(f" | {proj['role']}")
        
        if proj.get('tech_stack'):
            doc.add_paragraph(f"技术栈: {' | '.join(proj['tech_stack'])}")
        
        if proj.get('description'):
            doc.add_paragraph(proj['description'].strip())
        
        doc.add_paragraph()
    
    @staticmethod
    def _add_education_item(doc, edu: Dict):
        """添加教育背景条目"""
        if edu.get('school'):
            p = doc.add_paragraph()
            p.add_run(edu['school']).bold = True
            info = []
            if edu.get('degree'):
                info.append(edu['degree'])
            if edu.get('major'):
                info.append(edu['major'])
            if info:
                p.add_run(f" | {' '.join(info)}")
        
        if edu.get('start_date') or edu.get('end_date'):
            date = f"{format_date(edu.get('start_date', ''))} - {format_date(edu.get('end_date', ''))}"
            doc.add_paragraph(date)
        
        doc.add_paragraph()


class HTMLExporter:
    """HTML导出器"""
    
    @staticmethod
    def export(resume_data: Dict, output_path: str) -> bool:
        """导出为HTML文件"""
        try:
            html_content = HTMLExporter.generate(resume_data)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return True
            
        except Exception as e:
            print(f"HTML导出失败: {e}")
            return False
    
    @staticmethod
    def generate(resume_data: Dict) -> str:
        """生成HTML内容"""
        name = html.escape(resume_data.get('name', ''))
        job_title = html.escape(resume_data.get('job_title', ''))
        
        html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{name} - 简历</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: 'Microsoft YaHei', 'PingFang SC', sans-serif; line-height: 1.6; color: #333; background: #f5f5f5; }}
        .container {{ max-width: 800px; margin: 40px auto; background: white; box-shadow: 0 2px 20px rgba(0,0,0,0.1); }}
        .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 40px; text-align: center; }}
        .header h1 {{ font-size: 2.5em; margin-bottom: 10px; }}
        .header .job-title {{ font-size: 1.2em; opacity: 0.9; }}
        .contact {{ margin-top: 15px; font-size: 0.9em; }}
        .contact span {{ margin: 0 15px; }}
        .content {{ padding: 40px; }}
        .section {{ margin-bottom: 30px; }}
        .section h2 {{ color: #667eea; border-bottom: 2px solid #667eea; padding-bottom: 10px; margin-bottom: 20px; font-size: 1.3em; }}
        .section p {{ margin-bottom: 10px; text-align: justify; }}
        .skills {{ display: flex; flex-wrap: wrap; gap: 8px; }}
        .skill-tag {{ background: #f0f0f0; padding: 5px 12px; border-radius: 15px; font-size: 0.9em; }}
        .item {{ margin-bottom: 20px; }}
        .item-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px; }}
        .item-title {{ font-weight: bold; color: #333; }}
        .item-date {{ color: #888; font-size: 0.9em; }}
        .item-subtitle {{ color: #666; margin-bottom: 8px; }}
        .item-content {{ color: #555; }}
        .achievement {{ color: #666; font-size: 0.95em; margin-left: 20px; position: relative; }}
        .achievement::before {{ content: '◆'; position: absolute; left: -15px; color: #667eea; }}
        @media print {{ body {{ background: white; }} .container {{ box-shadow: none; margin: 0; }} }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>{name}</h1>
            <p class="job-title">{job_title}</p>
            <div class="contact">
                <span>📱 {resume_data.get('phone', '')}</span>
                <span>✉️ {resume_data.get('email', '')}</span>
                <span>📍 {resume_data.get('location', '')}</span>
            </div>
        </div>
        <div class="content">
"""
        
        if resume_data.get('self_introduction'):
            html_content += f"""
            <div class="section">
                <h2>个人简介</h2>
                <p>{html.escape(resume_data['self_introduction'])}</p>
            </div>
"""
        
        if resume_data.get('skills'):
            skills_html = ''.join([f'<span class="skill-tag">{html.escape(s)}</span>' for s in resume_data['skills']])
            html_content += f"""
            <div class="section">
                <h2>专业技能</h2>
                <div class="skills">{skills_html}</div>
            </div>
"""
        
        if resume_data.get('work_experience'):
            work_html = ''
            for exp in resume_data['work_experience']:
                work_html += f"""
            <div class="item">
                <div class="item-header">
                    <span class="item-title">{html.escape(exp.get('company', ''))}</span>
                    <span class="item-date">{format_date(exp.get('start_date', ''))} - {format_date(exp.get('end_date', ''))}</span>
                </div>
                <p class="item-subtitle">{html.escape(exp.get('position', ''))}</p>
                <p class="item-content">{html.escape(exp.get('description', ''))}</p>
"""
                if exp.get('achievements'):
                    for ach in exp['achievements'][:3]:
                        work_html += f'<p class="achievement">{html.escape(ach)}</p>'
                work_html += '</div>'
            
            html_content += f"""
            <div class="section">
                <h2>工作经历</h2>
                {work_html}
            </div>
"""
        
        if resume_data.get('project_experience'):
            proj_html = ''
            for proj in resume_data['project_experience']:
                proj_html += f"""
            <div class="item">
                <div class="item-header">
                    <span class="item-title">{html.escape(proj.get('name', ''))}</span>
                    <span class="item-date">{format_date(proj.get('start_date', ''))} - {format_date(proj.get('end_date', ''))}</span>
                </div>
                <p class="item-subtitle">{html.escape(proj.get('role', ''))}</p>
                <p class="item-content">{html.escape(proj.get('description', ''))}</p>
            </div>
"""
            
            html_content += f"""
            <div class="section">
                <h2>项目经历</h2>
                {proj_html}
            </div>
"""
        
        if resume_data.get('education'):
            edu_html = ''
            for edu in resume_data['education']:
                edu_html += f"""
            <div class="item">
                <div class="item-header">
                    <span class="item-title">{html.escape(edu.get('school', ''))}</span>
                    <span class="item-date">{format_date(edu.get('start_date', ''))} - {format_date(edu.get('end_date', ''))}</span>
                </div>
                <p class="item-subtitle">{html.escape(edu.get('degree', ''))} {html.escape(edu.get('major', ''))}</p>
            </div>
"""
            
            html_content += f"""
            <div class="section">
                <h2>教育背景</h2>
                {edu_html}
            </div>
"""
        
        if resume_data.get('certificates'):
            cert_html = ''.join([f'<p>• {html.escape(c)}</p>' for c in resume_data['certificates']])
            html_content += f"""
            <div class="section">
                <h2>证书资质</h2>
                {cert_html}
            </div>
"""
        
        if resume_data.get('awards'):
            award_html = ''.join([f'<p>• {html.escape(a)}</p>' for a in resume_data['awards']])
            html_content += f"""
            <div class="section">
                <h2>荣誉奖励</h2>
                {award_html}
            </div>
"""
        
        html_content += """
        </div>
    </div>
</body>
</html>
"""
        
        return html_content


def export_to_word(resume_data: Dict, output_path: str) -> bool:
    """导出简历为Word文档"""
    return WordExporter.export(resume_data, output_path)


def export_to_html(resume_data: Dict, output_path: str) -> bool:
    """导出简历为HTML文件"""
    return HTMLExporter.export(resume_data, output_path)


def get_available_export_formats() -> List[Dict]:
    """获取可用的导出格式"""
    return [
        {
            'id': 'pdf',
            'name': 'PDF文档',
            'description': '标准PDF格式，适合打印和投递',
            'extensions': ['.pdf']
        },
        {
            'id': 'word',
            'name': 'Word文档',
            'description': 'Word格式(.docx)，方便编辑修改',
            'extensions': ['.docx']
        },
        {
            'id': 'html',
            'name': '网页格式',
            'description': 'HTML格式，可在浏览器中查看',
            'extensions': ['.html']
        }
    ]
